"""Document parsers for local and fetched content formats."""

from __future__ import annotations

import re
from html import unescape

from deeprhetor.domain.assessment import ParsedSourceMetadata
from deeprhetor.domain.sources import ParsedDocument, ParsedSegment, RawDocument

PARSER_VERSION = "1.0.0"

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)


def guess_media_type(filename: str | None, fallback: str = "application/octet-stream") -> str:
    if not filename:
        return fallback
    lower = filename.lower()
    mapping = {
        ".pdf": "application/pdf",
        ".html": "text/html",
        ".htm": "text/html",
        ".md": "text/markdown",
        ".markdown": "text/markdown",
        ".txt": "text/plain",
        ".text": "text/plain",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    for ext, media in mapping.items():
        if lower.endswith(ext):
            return media
    return fallback


def _segments_from_blocks(
    blocks: list[tuple[str, int | None, str | None]],
) -> tuple[str, list[ParsedSegment]]:
    """Build full text and char-offset segments from (text, page, section_path) blocks."""
    segments: list[ParsedSegment] = []
    parts: list[str] = []
    offset = 0
    for text, page, section in blocks:
        cleaned = text.strip()
        if not cleaned:
            continue
        if parts:
            offset += 2  # "\n\n" join
        char_start = offset
        char_end = offset + len(cleaned)
        segments.append(
            ParsedSegment(
                text=cleaned,
                page=page,
                section_path=section,
                char_start=char_start,
                char_end=char_end,
                status="pending",
            )
        )
        parts.append(cleaned)
        offset = char_end
    return "\n\n".join(parts), segments


class PlainTextParser:
    supported = frozenset({"text/plain", "text/markdown"})

    def supports(self, media_type: str) -> bool:
        return media_type.split(";", 1)[0].strip().lower() in self.supported

    async def parse(self, document: RawDocument) -> ParsedDocument:
        text = document.content.decode("utf-8", errors="replace")
        media = document.media_type.split(";", 1)[0].strip().lower()
        if media == "text/markdown" or (document.filename or "").lower().endswith(
            (".md", ".markdown")
        ):
            return self._parse_markdown(document, text)
        return self._parse_plain(document, text)

    def _parse_plain(self, document: RawDocument, text: str) -> ParsedDocument:
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        if not paragraphs and text.strip():
            paragraphs = [text.strip()]
        full, segments = _segments_from_blocks([(p, None, None) for p in paragraphs])
        return ParsedDocument(
            media_type="text/plain",
            title=document.title or document.filename,
            text=full,
            segments=segments,
            parser="plaintext",
            parser_version=PARSER_VERSION,
            source_metadata=ParsedSourceMetadata(title=document.title or document.filename),
        )

    def _parse_markdown(self, document: RawDocument, text: str) -> ParsedDocument:
        blocks: list[tuple[str, int | None, str | None]] = []
        current_section: str | None = None
        buffer: list[str] = []

        def flush() -> None:
            nonlocal buffer
            joined = "\n".join(buffer).strip()
            if joined:
                blocks.append((joined, None, current_section))
            buffer = []

        for line in text.splitlines():
            match = _HEADING_RE.match(line)
            if match:
                flush()
                current_section = match.group(2).strip()
                blocks.append((line.strip(), None, current_section))
            else:
                if not line.strip() and buffer:
                    flush()
                elif line.strip() or buffer:
                    buffer.append(line)
        flush()
        if not blocks and text.strip():
            blocks.append((text.strip(), None, None))
        full, segments = _segments_from_blocks(blocks)
        title = document.title
        if not title:
            for seg in segments:
                if seg.section_path and seg.text.startswith("#"):
                    title = seg.section_path
                    break
            title = title or document.filename
        return ParsedDocument(
            media_type="text/markdown",
            title=title,
            text=full,
            segments=segments,
            parser="markdown",
            parser_version=PARSER_VERSION,
            source_metadata=ParsedSourceMetadata(title=title),
        )


class HtmlParser:
    supported = frozenset({"text/html", "application/xhtml+xml"})

    def supports(self, media_type: str) -> bool:
        return media_type.split(";", 1)[0].strip().lower() in self.supported

    async def parse(self, document: RawDocument) -> ParsedDocument:
        raw = document.content.decode("utf-8", errors="replace")
        title: str | None = document.title
        extracted = ""
        try:
            import trafilatura

            extracted = trafilatura.extract(raw, include_comments=False, include_tables=True) or ""
            meta = trafilatura.extract_metadata(raw)
            if meta and meta.title and not title:
                title = meta.title
        except Exception:
            extracted = ""

        if not extracted:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(raw, "html.parser")
            if not title and soup.title and soup.title.string:
                title = soup.title.string.strip()
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()
            extracted = soup.get_text("\n")

        # Split into paragraph-ish blocks
        blocks_text = [p.strip() for p in re.split(r"\n\s*\n", extracted) if p.strip()]
        if not blocks_text and extracted.strip():
            blocks_text = [extracted.strip()]
        # Detect simple markdown-like headings left by trafilatura
        blocks: list[tuple[str, int | None, str | None]] = []
        section: str | None = None
        for block in blocks_text:
            if len(block) < 120 and not block.endswith("."):
                section = block
            blocks.append((unescape(block), None, section))
        full, segments = _segments_from_blocks(blocks)
        return ParsedDocument(
            media_type="text/html",
            title=title or document.filename,
            text=full,
            segments=segments,
            parser="html",
            parser_version=PARSER_VERSION,
            source_metadata=ParsedSourceMetadata(
                title=title or document.filename,
                publisher_or_site=None,
            ),
        )


class PdfParser:
    """PDF text extraction: pymupdf first, OCR fallback for scanned/image pages."""

    supported = frozenset({"application/pdf"})
    # Pages below this character count are candidates for OCR.
    ocr_char_threshold: int = 40

    def supports(self, media_type: str) -> bool:
        return media_type.split(";", 1)[0].strip().lower() in self.supported

    async def parse(self, document: RawDocument) -> ParsedDocument:
        import fitz  # pymupdf

        blocks: list[tuple[str, int | None, str | None]] = []
        title = document.title
        warnings: list[str] = []
        used_ocr = False
        parser_name = "pymupdf"

        with fitz.open(stream=document.content, filetype="pdf") as pdf:
            meta = pdf.metadata or {}
            if not title:
                title = meta.get("title") or document.filename
            for page_index, page in enumerate(pdf, start=1):
                text = page.get_text("text").strip()
                if len(text) >= self.ocr_char_threshold:
                    for para in re.split(r"\n\s*\n", text):
                        cleaned = para.strip()
                        if cleaned:
                            blocks.append((cleaned, page_index, f"page/{page_index}"))
                    continue

                ocr_text, ocr_warning = _ocr_page(page)
                if ocr_warning:
                    warnings.append(f"page/{page_index}: {ocr_warning}")
                if ocr_text:
                    used_ocr = True
                    for para in re.split(r"\n\s*\n", ocr_text):
                        cleaned = para.strip()
                        if cleaned:
                            blocks.append((cleaned, page_index, f"page/{page_index}"))
                elif text:
                    # Keep sparse native text if OCR unavailable / empty.
                    blocks.append((text, page_index, f"page/{page_index}"))

        if used_ocr:
            parser_name = "pymupdf+ocr"
        full, segments = _segments_from_blocks(blocks)
        source_meta = ParsedSourceMetadata(title=title)
        if warnings:
            source_meta.extra = {"extraction_warnings": warnings}
        return ParsedDocument(
            media_type="application/pdf",
            title=title,
            text=full,
            segments=segments,
            parser=parser_name,
            parser_version=PARSER_VERSION,
            source_metadata=source_meta,
        )


def _ocr_page(page: object) -> tuple[str, str | None]:
    """Render a PDF page and OCR it. Gracefully skip if tesseract is missing."""
    try:
        import fitz  # pymupdf
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        return "", f"ocr_skipped_import: {exc}"

    try:
        # pix map at moderate DPI for OCR
        assert isinstance(page, fitz.Page)
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        text = pytesseract.image_to_string(image) or ""
        return text.strip(), None
    except pytesseract.TesseractNotFoundError:
        return "", "ocr_skipped_tesseract_missing"
    except Exception as exc:  # noqa: BLE001 — OCR is best-effort
        return "", f"ocr_failed: {exc}"


class DocxParser:
    supported = frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )

    def supports(self, media_type: str) -> bool:
        return media_type.split(";", 1)[0].strip().lower() in self.supported

    async def parse(self, document: RawDocument) -> ParsedDocument:
        from io import BytesIO

        from docx import Document as DocxDocument

        doc = DocxDocument(BytesIO(document.content))
        title = document.title or document.filename
        blocks: list[tuple[str, int | None, str | None]] = []
        section: str | None = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name if para.style is not None else "") or ""
            if style.startswith("Heading"):
                section = text
                blocks.append((text, None, section))
            else:
                blocks.append((text, None, section))
        full, segments = _segments_from_blocks(blocks)
        return ParsedDocument(
            media_type=next(iter(self.supported)),
            title=title,
            text=full,
            segments=segments,
            parser="python-docx",
            parser_version=PARSER_VERSION,
            source_metadata=ParsedSourceMetadata(title=title),
        )


class ParserRegistry:
    """Choose a DocumentParser by media type."""

    def __init__(self, parsers: list | None = None) -> None:
        self._parsers = parsers or [
            PdfParser(),
            HtmlParser(),
            DocxParser(),
            PlainTextParser(),
        ]

    def get(self, media_type: str) -> object:
        normalized = media_type.split(";", 1)[0].strip().lower()
        for parser in self._parsers:
            if parser.supports(normalized):
                return parser
        raise ValueError(f"no parser for media type: {media_type}")

    def supports(self, media_type: str) -> bool:
        try:
            self.get(media_type)
            return True
        except ValueError:
            return False

    async def parse(self, document: RawDocument) -> ParsedDocument:
        parser = self.get(document.media_type)
        return await parser.parse(document)  # type: ignore[no-any-return]


default_parser_registry = ParserRegistry()
